
#  python3 -m venv myenv

#  source myenv/bin/activate

#  pip install python-docx

#  python3 /Users/hidayetposcu/eslima/yardimci_dosyalar-2/word.py




import os
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# İşlem yapılacak klasör yolu
folder_path = "/Users/hidayetposcu/Eslima/yardimci_dosyalar-1"  # Kendi klasör yolunuzu yazın
new_font_name = "Arial"
table_background_color = "FFFFFF"  # Tablonun tamamı için beyaz zemin rengi (Hex format)

def update_tables_in_docx(file_path):
    doc = Document(file_path)

    for table in doc.tables:
        # Tablodaki tüm hücreler için işlem
        for row in table.rows:
            for cell in row.cells:
                # Hücrenin zemin rengini beyaz yap
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), table_background_color))
                )
                # Paragraf içerikleri üzerinde işlem
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = new_font_name
                        run.font.size = 80000  # Yazı boyutunu ayarlamak
                        run.font.italic = False  # İtalik iptal

    # Tablo dışındaki paragraflar için italik iptali
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.italic = False
            run.font.name = new_font_name

    # Değişiklikleri kaydet
    doc.save(file_path)

# Klasördeki tüm .docx dosyalarını işleme
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        print(f"Processing: {filename}")
        update_tables_in_docx(file_path)

print("Tüm dosyalar başarıyla güncellendi!")
